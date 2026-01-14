#!/usr/bin/env python3
"""
SRS Document Generator for Travel Planning & Recommendation System
Generates a comprehensive Software Requirements Specification document following IEEE 830-1998 standard
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import sys

def add_title_page(doc):
    """Add title page with project information"""
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Travel Planning & Recommendation System")
    run.font.size = Pt(24)
    run.font.bold = True
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("\nSoftware Requirements Specification")
    run.font.size = Pt(18)
    run.font.bold = True
    
    # Version
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run(f"\n\nVersion 1.0\n{datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(14)
    
    # Team info
    team = doc.add_paragraph()
    team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = team.add_run("\n\nPrepared by:\nDevelopment Team\nTravel Planning System Project")
    run.font.size = Pt(12)
    
    doc.add_page_break()

def add_toc(doc):
    """Add table of contents"""
    doc.add_heading("TABLE OF CONTENTS", 0)
    doc.add_paragraph("1.0 INTRODUCTION", style='List Bullet')
    doc.add_paragraph("1.1 Purpose", style='List Bullet 2')
    doc.add_paragraph("1.2 Scope of Project", style='List Bullet 2')
    doc.add_paragraph("1.3 Glossary", style='List Bullet 2')
    doc.add_paragraph("1.4 References", style='List Bullet 2')
    doc.add_paragraph("1.5 Overview of Document", style='List Bullet 2')
    doc.add_paragraph("1.6 Stakeholders", style='List Bullet 2')
    doc.add_paragraph("2.0 OVERALL DESCRIPTION", style='List Bullet')
    doc.add_paragraph("2.1 System Environment", style='List Bullet 2')
    doc.add_paragraph("2.2 Functional Requirements Specification", style='List Bullet 2')
    doc.add_paragraph("2.3 User Characteristics", style='List Bullet 2')
    doc.add_paragraph("2.4 Non-Functional Requirements", style='List Bullet 2')
    doc.add_paragraph("2.5 Budget (Estimated)", style='List Bullet 2')
    doc.add_paragraph("2.6 Feasibility Study", style='List Bullet 2')
    doc.add_paragraph("3.0 REQUIREMENTS SPECIFICATION", style='List Bullet')
    doc.add_paragraph("3.1 External Interface Requirements", style='List Bullet 2')
    doc.add_paragraph("3.2 Functional Requirements", style='List Bullet 2')
    doc.add_paragraph("3.3 Detailed Non-Functional Requirements", style='List Bullet 2')
    doc.add_paragraph("3.4 Data Model", style='List Bullet 2')
    doc.add_paragraph("APPENDICES", style='List Bullet')
    doc.add_page_break()

def add_list_of_figures(doc):
    """Add list of figures"""
    doc.add_heading("LIST OF FIGURES", 0)
    figures = [
        "Figure 1: System Use Case Diagram",
        "Figure 2: User Role Use Case Diagram",
        "Figure 3: Admin Role Use Case Diagram",
        "Figure 4: Data Model Class Diagram",
        "Figure 5: Travel Plan State Diagram",
        "Figure 6: Itinerary Generation Activity Diagram"
    ]
    for fig in figures:
        doc.add_paragraph(fig, style='List Bullet')
    doc.add_page_break()

def add_introduction(doc):
    """Add Section 1.0 INTRODUCTION"""
    doc.add_heading("1.0 INTRODUCTION", 1)
    
    # 1.1 Purpose
    doc.add_heading("1.1 Purpose", 2)
    doc.add_paragraph(
        "This Software Requirements Specification (SRS) document provides a complete description of the "
        "Travel Planning & Recommendation System. It defines the functional and non-functional requirements, "
        "system constraints, and interfaces for the system."
    )
    doc.add_paragraph("Intended Audience:")
    audiences = [
        "Developers: For implementation guidance and technical specifications",
        "Testers: For test case development and validation criteria",
        "Project Managers: For project planning and resource allocation",
        "Stakeholders: For understanding system capabilities and limitations"
    ]
    for audience in audiences:
        doc.add_paragraph(audience, style='List Bullet')
    
    # 1.2 Scope of Project
    doc.add_heading("1.2 Scope of Project", 2)
    doc.add_paragraph(
        "Project Name: Travel Planning & Recommendation System"
    )
    doc.add_paragraph(
        "The system is a rule-based travel planning and recommendation platform that helps users plan trips "
        "by providing personalized destination recommendations, hotel selections, and automated itinerary generation."
    )
    doc.add_paragraph("Main Features:")
    features = [
        "User preference management (budget, interests, travel party size)",
        "Destination recommendations based on rule-based algorithms",
        "Hotel recommendations matching budget constraints",
        "Transport option suggestions based on distance",
        "Automated itinerary generation",
        "Travel plan creation and management"
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_paragraph("Goals and Benefits:")
    benefits = [
        "Simplify travel planning process",
        "Provide personalized recommendations without AI/ML complexity",
        "Reduce time spent researching destinations and accommodations",
        "Generate structured travel itineraries automatically"
    ]
    for benefit in benefits:
        doc.add_paragraph(benefit, style='List Bullet')
    
    doc.add_paragraph("User Roles Supported:")
    roles = [
        "Guest User: Browse destinations and view information",
        "Registered User: Set preferences, create travel plans, generate itineraries",
        "Administrator: Manage destinations, hotels, and transport options"
    ]
    for role in roles:
        doc.add_paragraph(role, style='List Bullet')
    
    # 1.3 Glossary
    doc.add_heading("1.3 Glossary", 2)
    glossary = {
        "API": "Application Programming Interface",
        "CORS": "Cross-Origin Resource Sharing",
        "CRUD": "Create, Read, Update, Delete operations",
        "Django": "Python web framework used for backend development",
        "DRF": "Django REST Framework",
        "JWT": "JSON Web Token for authentication",
        "MySQL": "Relational database management system",
        "React": "JavaScript library for building user interfaces",
        "REST": "Representational State Transfer",
        "SPA": "Single Page Application",
        "SRS": "Software Requirements Specification"
    }
    for term, definition in glossary.items():
        p = doc.add_paragraph()
        p.add_run(f"{term}: ").bold = True
        p.add_run(definition)
    
    # 1.4 References
    doc.add_heading("1.4 References", 2)
    references = [
        "IEEE Std 830-1998: IEEE Recommended Practice for Software Requirements Specifications",
        "Django Documentation: https://docs.djangoproject.com/",
        "React Documentation: https://react.dev/",
        "Django REST Framework: https://www.django-rest-framework.org/",
        "MySQL Documentation: https://dev.mysql.com/doc/"
    ]
    for ref in references:
        doc.add_paragraph(ref, style='List Bullet')
    
    # 1.5 Overview of Document
    doc.add_heading("1.5 Overview of Document", 2)
    doc.add_paragraph(
        "This document is organized into three main sections:"
    )
    doc.add_paragraph(
        "Section 1 (Introduction) provides an overview of the document, project scope, and terminology."
    )
    doc.add_paragraph(
        "Section 2 (Overall Description) describes the system environment, functional requirements, "
        "user characteristics, and feasibility analysis."
    )
    doc.add_paragraph(
        "Section 3 (Requirements Specification) provides detailed functional and non-functional requirements, "
        "interface specifications, and data models."
    )
    
    # 1.6 Stakeholders
    doc.add_heading("1.6 Stakeholders", 2)
    doc.add_paragraph("Primary Stakeholders:")
    primary = [
        "Project Lead: Responsible for overall project direction and delivery",
        "Development Team: Backend and frontend developers implementing the system",
        "Database Administrator: Managing MySQL database and data integrity"
    ]
    for stakeholder in primary:
        doc.add_paragraph(stakeholder, style='List Bullet')
    
    doc.add_paragraph("Secondary Stakeholders:")
    secondary = [
        "End Users: Travelers seeking planning assistance",
        "Travel Service Providers: Hotels and transportation companies",
        "System Administrators: Maintaining system operations"
    ]
    for stakeholder in secondary:
        doc.add_paragraph(stakeholder, style='List Bullet')
    
    doc.add_page_break()

def add_overall_description(doc):
    """Add Section 2.0 OVERALL DESCRIPTION"""
    doc.add_heading("2.0 OVERALL DESCRIPTION", 1)
    
    # 2.1 System Environment
    doc.add_heading("2.1 System Environment", 2)
    doc.add_paragraph("Architecture: Client-Server with RESTful API")
    doc.add_paragraph("Technology Stack:")
    doc.add_paragraph("Frontend:", style='List Bullet')
    doc.add_paragraph("React 19", style='List Bullet 2')
    doc.add_paragraph("JavaScript/ES6+", style='List Bullet 2')
    doc.add_paragraph("CSS3", style='List Bullet 2')
    doc.add_paragraph("Axios for API calls", style='List Bullet 2')
    
    doc.add_paragraph("Backend:", style='List Bullet')
    doc.add_paragraph("Django 5.0", style='List Bullet 2')
    doc.add_paragraph("Django REST Framework 3.14", style='List Bullet 2')
    doc.add_paragraph("Python 3.10+", style='List Bullet 2')
    
    doc.add_paragraph("Database:", style='List Bullet')
    doc.add_paragraph("MySQL 8.0", style='List Bullet 2')
    
    # Add PlantUML diagram
    doc.add_paragraph("\nSystem Use Case Diagram (PlantUML):")
    plantuml_code = """@startuml
actor "User" as user
actor "Admin" as admin

rectangle "Travel Planning System" {
  usecase "Browse Destinations" as UC1
  usecase "Set Preferences" as UC2
  usecase "Create Travel Plan" as UC3
  usecase "Generate Itinerary" as UC4
  usecase "Manage Destinations" as UC5
  usecase "Manage Hotels" as UC6
  usecase "Manage Transport" as UC7
}

user --> UC1
user --> UC2
user --> UC3
user --> UC4

admin --> UC5
admin --> UC6
admin --> UC7

@enduml"""
    
    p = doc.add_paragraph(plantuml_code)
    p.style = 'Intense Quote'
    
    # 2.2 Functional Requirements Specification
    doc.add_heading("2.2 Functional Requirements Specification", 2)
    
    doc.add_paragraph("User Role: Registered User")
    doc.add_paragraph("\nUse Case 1: Set Travel Preferences")
    doc.add_paragraph("Description: User defines budget level, interests, and number of travelers")
    doc.add_paragraph("User Story: As a registered user, I want to set my travel preferences so that I receive personalized recommendations")
    doc.add_paragraph("Acceptance Criteria:")
    criteria = [
        "User can select budget level (low, medium, high)",
        "User can select interest category (beach, wildlife, historical, city tour, adventure, culture)",
        "User can specify number of travelers",
        "Preferences are saved and can be updated",
        "System validates all inputs before saving"
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet')
    
    doc.add_paragraph("\nUse Case 2: Get Destination Recommendations")
    doc.add_paragraph("Description: System recommends destinations based on user preferences")
    doc.add_paragraph("User Story: As a registered user, I want to receive destination recommendations so that I can choose where to travel")
    doc.add_paragraph("Acceptance Criteria:")
    criteria = [
        "System filters destinations by budget level",
        "System filters destinations by interest category",
        "Results display destination name, description, and image",
        "User can view detailed information for each destination",
        "Recommendations update when preferences change"
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet')
    
    doc.add_paragraph("\nUse Case 3: Create Travel Plan")
    doc.add_paragraph("Description: User creates a complete travel plan with destination, hotel, and dates")
    doc.add_paragraph("User Story: As a registered user, I want to create a travel plan so that I can organize my trip")
    doc.add_paragraph("Acceptance Criteria:")
    criteria = [
        "User can select destination from recommendations",
        "User can choose hotel based on budget",
        "User can specify travel dates",
        "System calculates total budget",
        "Travel plan is saved to user's account"
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet')
    
    doc.add_paragraph("\nUse Case 4: Generate Itinerary")
    doc.add_paragraph("Description: System automatically generates day-by-day itinerary")
    doc.add_paragraph("User Story: As a registered user, I want an automated itinerary so that I have a structured plan for my trip")
    doc.add_paragraph("Acceptance Criteria:")
    criteria = [
        "Day 1 includes arrival and hotel check-in",
        "Middle days include activities and attractions",
        "Last day includes shopping and return",
        "Each day has time-based schedule",
        "Itinerary can be viewed and printed"
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet')
    
    # Admin use cases
    doc.add_paragraph("\nUser Role: Administrator")
    doc.add_paragraph("\nUse Case 5: Manage Destinations")
    doc.add_paragraph("Description: Admin adds, updates, or removes destinations")
    doc.add_paragraph("User Story: As an administrator, I want to manage destinations so that users have up-to-date travel options")
    doc.add_paragraph("Acceptance Criteria:")
    criteria = [
        "Admin can add new destinations with all details",
        "Admin can update existing destination information",
        "Admin can deactivate destinations",
        "Changes reflect immediately in user recommendations",
        "System validates required fields"
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet')
    
    # 2.3 User Characteristics
    doc.add_heading("2.3 User Characteristics", 2)
    doc.add_paragraph("Registered User Profile:")
    doc.add_paragraph("Technical Proficiency: Basic to intermediate computer skills")
    doc.add_paragraph("Domain Knowledge: General understanding of travel planning")
    doc.add_paragraph("Frequency of Use: Occasional (planning trips)")
    doc.add_paragraph("Special Needs: Mobile-responsive interface for on-the-go access")
    
    doc.add_paragraph("\nAdministrator Profile:")
    doc.add_paragraph("Technical Proficiency: Advanced computer and database skills")
    doc.add_paragraph("Domain Knowledge: Travel industry expertise")
    doc.add_paragraph("Frequency of Use: Regular (daily management)")
    doc.add_paragraph("Special Needs: Bulk data management capabilities")
    
    # 2.4 Non-Functional Requirements
    doc.add_heading("2.4 Non-Functional Requirements", 2)
    nfr = [
        "Performance: Page load time < 3 seconds, API response time < 1 second",
        "Security: Password hashing, HTTPS encryption, input validation",
        "Usability: Intuitive interface, mobile-responsive design",
        "Reliability: 99% uptime, automatic error recovery",
        "Scalability: Support 1000+ concurrent users",
        "Maintainability: Modular code, comprehensive documentation"
    ]
    for req in nfr:
        doc.add_paragraph(req, style='List Bullet')
    
    # 2.5 Budget (Estimated)
    doc.add_heading("2.5 Budget (Estimated)", 2)
    doc.add_paragraph("Budget Breakdown (Student Project with GitHub Student Benefits):")
    
    budget_items = [
        ("Development Costs", "$0 (self-developed)"),
        ("Design (UI/UX)", "$0 (self-designed)"),
        ("Domain Name", "$0 (free via GitHub Student Pack)"),
        ("Hosting (DigitalOcean)", "$0 (free credits via GitHub Student Pack)"),
        ("Database Hosting", "$0 (free tier MySQL)"),
        ("SSL Certificate", "$0 (Let's Encrypt)"),
        ("Development Tools", "$0 (free via GitHub Student Pack - JetBrains, VS Code)"),
        ("Testing Tools", "$0 (open source)"),
        ("Marketing", "$0 (organic/social media)"),
        ("Maintenance (Year 1)", "$0 (self-maintained)"),
        ("Contingency (10%)", "$0")
    ]
    
    for item, cost in budget_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{item}: ").bold = True
        p.add_run(cost)
    
    p = doc.add_paragraph()
    p.add_run("Total Estimated Budget: ").bold = True
    p.add_run("$0 (leveraging free student resources)")
    
    doc.add_paragraph("\nNote: This budget assumes utilization of GitHub Student Developer Pack benefits including:")
    benefits = [
        "Free domain name for 1 year",
        "DigitalOcean $200 credit",
        "Azure $100 credit",
        "Free JetBrains IDE licenses",
        "Free GitHub Pro",
        "Free hosting credits from multiple providers"
    ]
    for benefit in benefits:
        doc.add_paragraph(benefit, style='List Bullet')
    
    # 2.6 Feasibility Study
    doc.add_heading("2.6 Feasibility Study", 2)
    
    doc.add_heading("2.6.1 Technical Feasibility", 3)
    doc.add_paragraph("Analysis:")
    doc.add_paragraph("Technology Availability: All required technologies (Django, React, MySQL) are mature, well-documented, and freely available")
    doc.add_paragraph("Team Expertise: Development team has experience with Python, JavaScript, and web development")
    doc.add_paragraph("Infrastructure: Standard LAMP stack requirements are well-supported")
    doc.add_paragraph("Scalability: Chosen architecture supports horizontal scaling")
    doc.add_paragraph("Integration: REST API pattern enables easy integration with external services")
    p = doc.add_paragraph()
    p.add_run("Conclusion: ").bold = True
    p.add_run("FEASIBLE - All technical requirements can be met with available technologies and skills")
    
    doc.add_heading("2.6.2 Economic Feasibility", 3)
    doc.add_paragraph("Analysis:")
    doc.add_paragraph("Initial Investment: $0 leveraging GitHub Student Pack")
    doc.add_paragraph("Recurring Costs: Minimal ($10-20/month after free credits expire)")
    doc.add_paragraph("Revenue Potential: Can be monetized through premium features, affiliate partnerships")
    doc.add_paragraph("Break-even: Low cost structure enables quick break-even")
    doc.add_paragraph("ROI: High potential return given minimal investment")
    p = doc.add_paragraph()
    p.add_run("Conclusion: ").bold = True
    p.add_run("FEASIBLE - Project is economically viable with minimal financial risk")
    
    doc.add_heading("2.6.3 Operational Feasibility", 3)
    doc.add_paragraph("Analysis:")
    doc.add_paragraph("User Acceptance: High likelihood - addresses real travel planning pain points")
    doc.add_paragraph("Market Demand: Growing online travel planning market")
    doc.add_paragraph("Training Requirements: Minimal - intuitive interface")
    doc.add_paragraph("Support Requirements: Basic FAQ and email support sufficient")
    doc.add_paragraph("Deployment: Standard web deployment process")
    doc.add_paragraph("Maintenance: Straightforward with modular architecture")
    p = doc.add_paragraph()
    p.add_run("Conclusion: ").bold = True
    p.add_run("FEASIBLE - System is operationally viable and user-friendly")
    
    doc.add_heading("2.6.4 Legal and Risk Feasibility", 3)
    doc.add_paragraph("Analysis:")
    doc.add_paragraph("Regulatory Compliance: Standard web application compliance (GDPR, data privacy)")
    doc.add_paragraph("Licensing: All open-source components with permissive licenses")
    doc.add_paragraph("Data Protection: User data handled per privacy best practices")
    doc.add_paragraph("Liability: Standard terms of service limit liability")
    doc.add_paragraph("Risk Mitigation: Regular backups, security updates, monitoring")
    p = doc.add_paragraph()
    p.add_run("Conclusion: ").bold = True
    p.add_run("FEASIBLE - No significant legal barriers or unmanageable risks")
    
    doc.add_heading("2.6.5 Overall Feasibility Conclusion", 3)
    p = doc.add_paragraph()
    p.add_run("Recommendation: ").bold = True
    p.add_run("GO - Proceed with development")
    doc.add_paragraph("\nKey Strengths:")
    strengths = [
        "Zero initial cost leveraging student benefits",
        "Mature, reliable technology stack",
        "Clear market need and use case",
        "Scalable architecture for future growth",
        "Manageable technical complexity"
    ]
    for strength in strengths:
        doc.add_paragraph(strength, style='List Bullet')
    
    doc.add_paragraph("\nPotential Challenges and Mitigation:")
    challenges = [
        "Challenge: Maintaining data accuracy → Mitigation: Admin dashboard for regular updates",
        "Challenge: User adoption → Mitigation: Focus on UX and social media marketing",
        "Challenge: Scaling costs → Mitigation: Optimize resources, implement caching"
    ]
    for challenge in challenges:
        doc.add_paragraph(challenge, style='List Bullet')
    
    doc.add_page_break()

def add_requirements_specification(doc):
    """Add Section 3.0 REQUIREMENTS SPECIFICATION"""
    doc.add_heading("3.0 REQUIREMENTS SPECIFICATION", 1)
    
    # 3.1 External Interface Requirements
    doc.add_heading("3.1 External Interface Requirements", 2)
    
    doc.add_heading("3.1.1 User Interfaces", 3)
    ui_components = [
        "Registration/Login screens with form validation",
        "Dashboard showing user preferences and travel plans",
        "Preference settings page with dropdown and radio selections",
        "Destination browse/search page with filtering",
        "Travel plan creation wizard (multi-step form)",
        "Itinerary view with day-by-day breakdown",
        "Admin panel for data management"
    ]
    for component in ui_components:
        doc.add_paragraph(component, style='List Bullet')
    
    doc.add_heading("3.1.2 Hardware Interfaces", 3)
    doc.add_paragraph("Client Side:")
    doc.add_paragraph("Any device with web browser (desktop, tablet, mobile)")
    doc.add_paragraph("Minimum screen resolution: 360x640 pixels")
    doc.add_paragraph("\nServer Side:")
    doc.add_paragraph("Linux server (Ubuntu 20.04+)")
    doc.add_paragraph("Minimum 2GB RAM, 20GB storage")
    doc.add_paragraph("MySQL database server")
    
    doc.add_heading("3.1.3 Software Interfaces", 3)
    interfaces = [
        "Django REST Framework: Provides RESTful API endpoints",
        "MySQL Database: Stores application data",
        "React Frontend: Single Page Application consuming API",
        "CORS Middleware: Enables cross-origin requests",
        "Django Admin: Built-in admin interface"
    ]
    for interface in interfaces:
        doc.add_paragraph(interface, style='List Bullet')
    
    doc.add_heading("3.1.4 Communication Interfaces", 3)
    doc.add_paragraph("Protocol: HTTP/HTTPS")
    doc.add_paragraph("Data Format: JSON")
    doc.add_paragraph("API Endpoints: RESTful URLs")
    doc.add_paragraph("Authentication: Session-based (can be extended to JWT)")
    doc.add_paragraph("Port: 8000 (backend), 3000 (frontend development)")
    
    # 3.2 Functional Requirements
    doc.add_heading("3.2 Functional Requirements", 2)
    
    requirements = [
        {
            "id": "FR-001",
            "title": "User Registration",
            "desc": "System shall allow new users to create accounts with username, email, and password",
            "input": "Username, email, password",
            "output": "User account created, confirmation message",
            "rules": "Password must be minimum 8 characters; Email must be unique"
        },
        {
            "id": "FR-002",
            "title": "User Authentication",
            "desc": "System shall authenticate users using username/email and password",
            "input": "Credentials (username/email, password)",
            "output": "Session token, user profile",
            "rules": "Max 5 failed attempts before temporary lockout"
        },
        {
            "id": "FR-003",
            "title": "Preference Management",
            "desc": "System shall allow users to set and update travel preferences",
            "input": "Budget level, interest category, number of travelers",
            "output": "Saved preferences, confirmation",
            "rules": "Budget: {low, medium, high}; Interest: {beach, wildlife, historical, city_tour, adventure, culture}"
        },
        {
            "id": "FR-004",
            "title": "Destination Recommendation",
            "desc": "System shall recommend destinations using rule-based filtering",
            "input": "User preferences (budget, interest)",
            "output": "List of matching destinations",
            "rules": "IF budget=low THEN filter budget_level=low; IF interest=beach THEN filter category=beach"
        },
        {
            "id": "FR-005",
            "title": "Hotel Recommendation",
            "desc": "System shall recommend hotels based on destination and budget",
            "input": "Destination ID, budget level",
            "output": "List of hotels with details",
            "rules": "IF budget=low THEN 1-2 star; IF budget=medium THEN 3 star; IF budget=high THEN 4-5 star"
        },
        {
            "id": "FR-006",
            "title": "Travel Plan Creation",
            "desc": "System shall allow users to create complete travel plans",
            "input": "Destination, hotel, dates, travelers",
            "output": "Created travel plan with total budget",
            "rules": "Start date must be future date; End date must be after start date"
        },
        {
            "id": "FR-007",
            "title": "Itinerary Generation",
            "desc": "System shall automatically generate day-by-day itineraries",
            "input": "Travel plan ID",
            "output": "Detailed itinerary with activities per day",
            "rules": "Day 1: Arrival; Middle days: Activities; Last day: Shopping and return"
        }
    ]
    
    for req in requirements:
        p = doc.add_paragraph()
        p.add_run(f"{req['id']}: {req['title']}").bold = True
        doc.add_paragraph(f"Description: {req['desc']}")
        doc.add_paragraph(f"Input: {req['input']}")
        doc.add_paragraph(f"Output: {req['output']}")
        doc.add_paragraph(f"Business Rules: {req['rules']}")
        doc.add_paragraph()
    
    # 3.3 Detailed Non-Functional Requirements
    doc.add_heading("3.3 Detailed Non-Functional Requirements", 2)
    
    doc.add_heading("3.3.1 Performance Requirements", 3)
    perf = [
        "Page load time: Maximum 3 seconds on 3G connection",
        "API response time: Maximum 1 second for 95% of requests",
        "Database query time: Maximum 500ms",
        "Concurrent users: Support minimum 1000 simultaneous users",
        "Throughput: Handle 100 API requests per second"
    ]
    for req in perf:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading("3.3.2 Security Requirements", 3)
    security = [
        "Password Storage: Use bcrypt hashing with salt",
        "Data Transmission: HTTPS/TLS 1.2 or higher",
        "Input Validation: Sanitize all user inputs to prevent SQL injection and XSS",
        "Authentication: Session-based with secure cookies",
        "Authorization: Role-based access control (User, Admin)",
        "Data Privacy: Comply with GDPR principles",
        "Backup: Daily automated backups with 30-day retention"
    ]
    for req in security:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading("3.3.3 Reliability Requirements", 3)
    reliability = [
        "Uptime: 99% availability (maximum 7 hours downtime per month)",
        "MTBF (Mean Time Between Failures): Minimum 720 hours",
        "MTTR (Mean Time To Repair): Maximum 2 hours",
        "Data Integrity: Database transactions with ACID properties",
        "Error Recovery: Automatic retry for failed API calls",
        "Fault Tolerance: Graceful degradation when services unavailable"
    ]
    for req in reliability:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading("3.3.4 Usability Requirements", 3)
    usability = [
        "Learning Curve: New users should complete first travel plan within 10 minutes",
        "Accessibility: WCAG 2.1 Level AA compliance",
        "Mobile Responsiveness: Support devices from 360px width",
        "Browser Support: Chrome, Firefox, Safari, Edge (latest 2 versions)",
        "Error Messages: Clear, actionable error messages",
        "Help Documentation: Inline help and FAQ section"
    ]
    for req in usability:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading("3.3.5 Maintainability Requirements", 3)
    maintainability = [
        "Code Style: Follow PEP 8 for Python, ESLint for JavaScript",
        "Documentation: Inline comments for complex logic, README for setup",
        "Modularity: Separate concerns (models, views, serializers)",
        "Version Control: Git with meaningful commit messages",
        "Testing: Unit test coverage minimum 70%",
        "Code Review: Pull request review before merging"
    ]
    for req in maintainability:
        doc.add_paragraph(req, style='List Bullet')
    
    # 3.4 Data Model
    doc.add_heading("3.4 Data Model", 2)
    doc.add_paragraph("Entity Relationship Diagram (PlantUML Class Diagram):")
    
    data_model = """@startuml
class User {
  + id: Integer
  + username: String
  + email: String
  + password: String (hashed)
  + created_at: DateTime
  --
  + register()
  + login()
}

class UserPreference {
  + id: Integer
  + user_id: Integer [FK]
  + budget: Enum(low, medium, high)
  + interest: Enum(beach, wildlife, historical, city_tour, adventure, culture)
  + num_travelers: Integer
  + created_at: DateTime
  --
  + save_preferences()
  + get_preferences()
}

class Destination {
  + id: Integer
  + name: String
  + country: String
  + city: String
  + category: String
  + budget_level: Enum(low, medium, high)
  + description: Text
  + image_url: String
  + booking_url: String
  --
  + get_by_category()
  + get_by_budget()
}

class Hotel {
  + id: Integer
  + destination_id: Integer [FK]
  + name: String
  + star_rating: Integer (1-5)
  + price_per_night: Decimal
  + budget_category: Enum(low, medium, high)
  + amenities: Text
  + location: String
  --
  + get_by_destination()
  + get_by_budget()
}

class Transport {
  + id: Integer
  + origin: String
  + destination: String
  + type: Enum(bus, train, flight, car)
  + distance_km: Integer
  + duration_hours: Decimal
  + price: Decimal
  --
  + get_by_distance()
  + calculate_price()
}

class TravelPlan {
  + id: Integer
  + user_id: Integer [FK]
  + destination_id: Integer [FK]
  + hotel_id: Integer [FK]
  + transport_id: Integer [FK]
  + start_date: Date
  + end_date: Date
  + num_travelers: Integer
  + total_budget: Decimal
  + created_at: DateTime
  --
  + create_plan()
  + calculate_total()
}

class Itinerary {
  + id: Integer
  + travel_plan_id: Integer [FK]
  + day_number: Integer
  + date: Date
  + activities: Text
  + accommodation: String
  + notes: Text
  --
  + generate()
  + get_by_plan()
}

User "1" -- "1" UserPreference : has
User "1" -- "*" TravelPlan : creates
TravelPlan "*" -- "1" Destination : includes
TravelPlan "*" -- "1" Hotel : books
TravelPlan "*" -- "0..1" Transport : uses
TravelPlan "1" -- "*" Itinerary : has
Destination "1" -- "*" Hotel : contains

@enduml"""
    
    p = doc.add_paragraph(data_model)
    p.style = 'Intense Quote'
    
    doc.add_paragraph("\nState Diagram for Travel Plan:")
    state_diagram = """@startuml
[*] --> Draft : Create Plan
Draft --> Confirmed : Confirm Details
Confirmed --> Booked : Make Bookings
Booked --> Active : Start Date Reached
Active --> Completed : End Date Reached
Booked --> Cancelled : Cancel Plan
Confirmed --> Cancelled : Cancel Plan
Draft --> Cancelled : Cancel Plan
Completed --> [*]
Cancelled --> [*]

note right of Draft
  User is still selecting
  destination and hotel
end note

note right of Confirmed
  All details finalized
  but not booked yet
end note

@enduml"""
    
    p = doc.add_paragraph(state_diagram)
    p.style = 'Intense Quote'
    
    doc.add_page_break()

def add_appendices(doc):
    """Add appendices"""
    doc.add_heading("APPENDIX A: PlantUML Diagram Instructions", 1)
    
    doc.add_paragraph("Rendering PlantUML Diagrams Online:")
    online_steps = [
        "Visit http://www.plantuml.com/plantuml/",
        "Copy the PlantUML code from this document",
        "Paste into the online editor",
        "Click 'Submit' to generate the diagram",
        "Download or save the generated image"
    ]
    for step in online_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_paragraph("\nLocal PlantUML Setup:")
    local_steps = [
        "Install Java Runtime Environment (JRE 8 or later)",
        "Download plantuml.jar from http://plantuml.com/download",
        "Install Graphviz from https://graphviz.org/download/",
        "Save PlantUML code to a .puml file",
        "Run: java -jar plantuml.jar diagram.puml",
        "Generated PNG/SVG will be created in the same directory"
    ]
    for step in local_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_paragraph("\nEditing Diagrams:")
    doc.add_paragraph(
        "To modify diagrams, edit the PlantUML code in this document and regenerate. "
        "PlantUML syntax is simple and well-documented at https://plantuml.com/"
    )
    
    doc.add_paragraph("\nIDE Integration:")
    ide_list = [
        "VS Code: Install 'PlantUML' extension",
        "IntelliJ IDEA: Install 'PlantUML integration' plugin",
        "Eclipse: Install PlantUML plugin from marketplace",
        "Atom: Install 'plantuml-viewer' package"
    ]
    for ide in ide_list:
        doc.add_paragraph(ide, style='List Bullet')

def generate_srs_document():
    """Main function to generate the complete SRS document"""
    print("Generating Software Requirements Specification document...")
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Add sections
    print("Adding title page...")
    add_title_page(doc)
    
    print("Adding table of contents...")
    add_toc(doc)
    
    print("Adding list of figures...")
    add_list_of_figures(doc)
    
    print("Adding introduction...")
    add_introduction(doc)
    
    print("Adding overall description...")
    add_overall_description(doc)
    
    print("Adding requirements specification...")
    add_requirements_specification(doc)
    
    print("Adding appendices...")
    add_appendices(doc)
    
    # Save document
    filename = f"Travel_Planning_System_SRS_v1.0_{datetime.now().strftime('%Y%m%d')}.docx"
    doc.save(filename)
    
    print(f"\n✓ SRS document generated successfully: {filename}")
    print(f"✓ Document size: {len(doc.paragraphs)} paragraphs")
    print(f"✓ Document includes:")
    print("  - Title page with project information")
    print("  - Complete table of contents")
    print("  - All IEEE 830-1998 required sections")
    print("  - PlantUML diagrams (use case, class, state)")
    print("  - Detailed functional and non-functional requirements")
    print("  - Comprehensive feasibility study")
    print("  - Budget breakdown with student benefits")
    print("  - Appendices with PlantUML instructions")
    
    return filename

if __name__ == "__main__":
    try:
        filename = generate_srs_document()
        sys.exit(0)
    except Exception as e:
        print(f"\n✗ Error generating SRS document: {str(e)}")
        print("\nPlease ensure python-docx is installed:")
        print("  pip install python-docx")
        sys.exit(1)
